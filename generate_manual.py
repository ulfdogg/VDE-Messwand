from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SCREENSHOTS = '/home/vde/VDE-Messwand/manual_screenshots'

doc = Document()

# ── Seitenränder ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.5)

# ── Farben ────────────────────────────────────────────────────────────────────
ROT  = RGBColor(0xC0, 0x00, 0x00)
GRAU = RGBColor(0x40, 0x40, 0x40)
HELL = RGBColor(0x60, 0x60, 0x60)
BLAU = RGBColor(0x1F, 0x49, 0x7D)

# ── Hilfsfunktionen ───────────────────────────────────────────────────────────

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, data in kwargs.items():
        tag = OxmlElement(f'w:{edge}')
        for key, value in data.items():
            tag.set(qn(f'w:{key}'), value)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def h1(text):
    """Kapitelüberschrift – wird ins Inhaltsverzeichnis aufgenommen."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    p.style = doc.styles['Heading 1']
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(17)
    run.font.color.rgb = ROT
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    p.style = doc.styles['Heading 2']
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = BLAU
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    p.style = doc.styles['Heading 3']
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GRAU
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAU
    return p

def bullet(bold_prefix, text=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if text:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10.5)
        rb.font.color.rgb = GRAU
        rt = p.add_run(text)
        rt.font.size = Pt(10.5)
        rt.font.color.rgb = GRAU
    else:
        run = p.add_run(bold_prefix)
        run.font.size = Pt(10.5)
        run.font.color.rgb = GRAU
    return p

def numbered(text, bold_prefix=''):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10.5)
        rb.font.color.rgb = GRAU
        rt = p.add_run(text)
        rt.font.size = Pt(10.5)
        rt.font.color.rgb = GRAU
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = GRAU
    return p

def hint_box(text, color_hex='F2F2F2', border_hex='C00000'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    shade_cell(cell, color_hex)
    set_cell_border(cell,
        top    = {'val': 'single', 'sz': '6',  'color': border_hex},
        bottom = {'val': 'single', 'sz': '6',  'color': border_hex},
        left   = {'val': 'single', 'sz': '12', 'color': border_hex},
        right  = {'val': 'none'}
    )
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = GRAU
    doc.add_paragraph()
    return table

def simple_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        shade_cell(cell, '1F497D')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row in enumerate(rows):
        bg = 'F7F7F7' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = GRAU
    doc.add_paragraph()
    return table

def page_break():
    doc.add_page_break()

def screenshot(filename, caption=None, width_cm=14.5):
    """Bild einbetten mit optionaler Bildunterschrift."""
    path = os.path.join(SCREENSHOTS, filename)
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_after = Pt(12)
        rc = pc.add_run(caption)
        rc.italic = True
        rc.font.size = Pt(9)
        rc.font.color.rgb = HELL

def add_toc():
    """Fügt ein Word-Inhaltsverzeichnis-Feld ein (wird beim Öffnen in Word aktualisiert)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    run4 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run4._r.append(fldChar3)

    return paragraph

# ══════════════════════════════════════════════════════════════════════════════
#  TITELSEITE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Cm(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('VDE-Messwand')
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = ROT

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Bedienungshandbuch')
run.font.size = Pt(20)
run.font.color.rgb = GRAU

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Prüfsystem für elektrische Anlagen nach DIN VDE 0100-600')
run.font.size = Pt(12)
run.font.color.rgb = HELL

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0  |  2025')
run.font.size = Pt(10)
run.font.color.rgb = HELL

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  INHALTSVERZEICHNIS
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
run = p.add_run('Inhaltsverzeichnis')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = ROT

doc.add_paragraph()
add_toc()

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  BEDIENER
# ══════════════════════════════════════════════════════════════════════════════

h1('Bediener')

body('Dieser Abschnitt richtet sich an alle Personen, die das System für Prüfungen und Übungen nutzen – also Prüfer, Lehrkräfte und Auszubildende. Ein Zugang zum Admin-Bereich ist hier nicht erforderlich.')

h2('Startseite')
body('Die Startseite zeigt vier Hauptbereiche:')

simple_table(
    ['Schaltfläche', 'Beschreibung', 'Zielgruppe'],
    [
        ['Prüfungsmodus',   'Automatisierter Prüfungsablauf mit zufälligen Fehlern und Timer',              'Prüfer / Auszubildende'],
        ['Manuelle Fehler', 'Gezielte manuelle Auswahl von Fehlern je Stromkreis',                          'Lehrkräfte'],
        ['Übungsmodus',     'Lernvideos und Messanleitungen für alle Messverfahren und Geräte',              'Auszubildende'],
        ['Admin',           'Konfiguration, Datenbank, Netzwerk (Passwort geschützt)',                      'Administratoren'],
    ]
)

hint_box('Hinweis: Die Wallbox-Umschalttaste erscheint auf der Startseite nur dann, wenn die Wallbox im Admin-Bereich als installiert markiert wurde.', 'EFF3FF', '1F497D')
screenshot('01_startseite.png', 'Startseite der VDE-Messwand')

# ── Prüfungsmodus ─────────────────────────────────────────────────────────────
h2('Prüfungsmodus')

body('Im Prüfungsmodus wählt das System automatisch und zufällig Fehler aus und aktiviert die entsprechenden Relais. Ein Timer läuft mit. So kann eine realistische VDE-Prüfung simuliert werden, bei der der Prüfling die Fehler mit Messgeräten selbst finden muss.')

h3('Ablauf')
numbered('Startseite öffnen und "Prüfungsmodus" antippen.')
numbered('Eine neue Prüfung starten – das System vergibt automatisch eine Prüfungsnummer (z.B. VDE-1).')
numbered('Die Relais werden zufällig aktiviert. Welche Fehler aktiv sind, wird dem Prüfer angezeigt, dem Prüfling nicht.')
numbered('Der Timer läuft. Der Prüfling führt die Messungen durch.')
numbered('Nach Abschluss "Prüfung beenden" drücken. Das System setzt alle Relais zurück und speichert die Ergebnisse.')

h3('Hinweise')
bullet('Prüfungsnummern:', ' Fortlaufend, Format VDE-1, VDE-2 usw.')
bullet('Timer:', ' 20 Minuten (läuft sichtbar mit)')
bullet('Gruppen:', ' Relais in einer Gruppe werden gemeinsam geschaltet und zählen als ein Fehler')
bullet('Datenbank:', ' Jede Prüfung wird automatisch mit Datum, Uhrzeit und aktivierten Fehlern gespeichert')

hint_box('Wichtig: Solange eine Prüfung läuft, sind die Relais aktiv. Erst nach "Prüfung beenden" werden alle Relais zurückgesetzt.', 'FFF0F0', 'C00000')
screenshot('02_pruefungsmodus.png', 'Prüfungsmodus – Übersicht und Timer')

# ── Manuelle Fehler ───────────────────────────────────────────────────────────
h2('Manuelle Fehler')

body('Im Modus "Manuelle Fehler" kann gezielt ausgewählt werden, welche Fehler in welchem Stromkreis aktiv sein sollen. Das eignet sich für gezielte Übungen zu einzelnen Messverfahren.')

h3('Ablauf')
numbered('Startseite öffnen und "Manuelle Fehler" antippen.')
numbered('Die Ansicht zeigt alle konfigurierten Stromkreise.')
numbered('Innerhalb eines Stromkreises die gewünschten Fehler (Relais) auswählen.')
numbered('Mehrfachauswahl möglich – auch aus verschiedenen Stromkreisen gleichzeitig.')
numbered('"Fehler aktivieren" drücken – die ausgewählten Relais schalten.')
numbered('Zum Zurücksetzen: "Alle Relais zurücksetzen" drücken.')

h3('Hinweise')
bullet('Gruppenrelais:', ' Wenn ein Relais einer Gruppe gewählt wird, schalten automatisch alle Relais der Gruppe mit.')
bullet('Wallbox-Stromkreis:', ' Erscheint nur, wenn die Wallbox auf der Startseite aktiviert ist.')
bullet('Kein Timer:', ' Im manuellen Modus läuft kein automatischer Timer.')
screenshot('03_manuelle_fehler.png', 'Manuelle Fehlerauswahl nach Stromkreisen')

# ── Übungsmodus ───────────────────────────────────────────────────────────────
h2('Übungsmodus')

body('Der Übungsmodus stellt Lernmaterial zu allen relevanten Messverfahren der DIN VDE 0100-600 bereit. Jede Karte öffnet einen erklärenden Text. Zusätzlich kann ein passendes Video abgespielt werden.')

h3('Aufbau der Seite')
simple_table(
    ['Abschnitt', 'Inhalte'],
    [
        ['Vorbereitung',             'Vor der Messung, Sichtprüfung, Sicherheit (Spannungsfreischalten)'],
        ['Spannungsfreie Messungen', 'Schutzleiter-Durchgängigkeit (R-Low), Isolationswiderstand (RISO)'],
        ['Messungen unter Spannung', 'Spannungsmessung, Schleifenimpedanz, RCD Auslösezeit, RCD Auslösestrom'],
        ['Abschluss',                'Funktionsprüfung, Dokumentation'],
        ['Messgeräte',               'Fluke, Gossen Metrawatt, Benning – Bedienung und Funktionen'],
        ['Messarten mit Relais',     'Spannungsfrei üben, Unter Spannung üben (mit Relais-Steuerung)'],
    ]
)

h3('Bedienung')
numbered('Gewünschte Karte antippen.')
numbered('Ein Fenster öffnet sich mit dem erklärenden Text zur Messung.')
numbered('"Video ansehen" drücken, um das zugehörige Lernvideo zu starten.')
numbered('"Video ausblenden" schließt den Player, der Text bleibt sichtbar.')
numbered('Fenster über das X oder durch Klick außerhalb schließen.')

h3('Unterseiten mit Relais-Steuerung')
body('Die Unterseiten "Spannungsfrei üben" und "Unter Spannung üben" ermöglichen es, direkt aus dem Übungsmodus passende Relais zu aktivieren. Welche Relais dabei geschaltet werden, wird im Admin-Bereich unter "Übungsmodus-Konfiguration" festgelegt.')
screenshot('04_uebungsmodus.png', 'Übungsmodus – Themenübersicht')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ADMIN
# ══════════════════════════════════════════════════════════════════════════════

h1('Administrator')

body('Dieser Abschnitt richtet sich an Administratoren, die das System einrichten, pflegen und konfigurieren. Der Admin-Bereich ist durch einen Code gesichert.')

# ── Verbindung ────────────────────────────────────────────────────────────────
h2('Verbindung zum System')

h3('Möglichkeit 1 – Hotspot (Standard, kabellos)')
body('Das Gerät spannt bei Bedarf einen eigenen WLAN-Hotspot auf. Verbinden Sie sich mit einem Tablet, Laptop oder Smartphone:')
bullet('Netzwerkname (SSID):', ' VDE-Messwand-1')
bullet('Passwort:', ' vde12345')
bullet('Adresse im Browser:', ' http://192.168.50.1')

hint_box('Tipp: Der Hotspot lässt sich im Admin-Bereich unter "Netzwerk" ein- und ausschalten.', 'FFF8E1', 'C07000')

h3('Möglichkeit 2 – Bestehendes WLAN')
body('Das System kann in ein vorhandenes WLAN eingebunden werden. Die aktuelle IP-Adresse ist im Admin-Bereich unter "Netzwerk" einsehbar.')
bullet('http://<IP-Adresse>')
bullet('http://<Hostname>.local  (z.B. http://vde-messwand.local)')

h3('Möglichkeit 3 – Direktes Display (Kiosk-Modus)')
body('Das System verfügt über ein angeschlossenes Display, auf dem die Oberfläche direkt im Vollbild-Browser läuft. Eine Netzwerkverbindung ist dafür nicht notwendig.')

# ── Anmeldung ─────────────────────────────────────────────────────────────────
h2('Anmeldung')

body('Der Admin-Bereich ist durch einen numerischen Code gesichert. Der Standardcode lautet 1234.')
numbered('Auf der Startseite "Admin" antippen.')
numbered('Den Admin-Code über das Nummernpad eingeben.')
numbered('Mit "Bestätigen" abschicken.')
numbered('Bei korrektem Code öffnet sich das Admin-Panel.')

hint_box('Sicherheitshinweis: Den Code nach der Ersteinrichtung ändern. Standardcode ist 1234.', 'FFF0F0', 'C00000')
screenshot('05_admin_login.png', 'Admin-Anmeldung per Nummernpad')

h2('Admin-Panel Übersicht')

simple_table(
    ['Bereich', 'Funktion'],
    [
        ['Datenbank',                 'Prüfungsprotokolle einsehen, exportieren und löschen'],
        ['Netzwerk',                  'WLAN, Hotspot und Verbindungsadressen verwalten'],
        ['Relais-Verwaltung',         'Alle 64 Relais konfigurieren (Gruppe, Name, Kategorie, Stromkreis)'],
        ['Konfiguration',             'Stromkreise und Messkategorien anlegen und bearbeiten'],
        ['Übungsmodus-Konfiguration', 'Relais-Zuordnung für die Trainingsseiten festlegen'],
        ['Einstellungen',             'Admin-Code ändern, Wallbox-Installation verwalten'],
        ['Relay Status Monitor',      'Live-Ansicht aller 64 Relais (ein/aus)'],
        ['Test-Durchlauf',            'Alle Relais automatisch der Reihe nach testen'],
        ['Alle Relais zurücksetzen',  'Sofortiges Abschalten aller aktiven Relais'],
        ['System herunterfahren',     'Betriebssystem sicher herunterfahren'],
    ]
)
screenshot('06_admin_panel.png', 'Admin-Panel – Übersicht aller Verwaltungsbereiche')

# ── Relais-Verwaltung ─────────────────────────────────────────────────────────
h2('Relais-Verwaltung')

body('Hier werden alle 64 Relais des Systems konfiguriert. Jedes Relais kann mit einem Namen, einer Kategorie, einem Stromkreis und einer Gruppe versehen werden.')

h3('Ansicht und Filter')
bullet('"Alle":', ' Alle 64 Relais anzeigen')
bullet('"Nur Gruppen":', ' Nur Relais, die einer Gruppe zugehören')
bullet('"Einzeln":', ' Nur Relais ohne Gruppe')
bullet('"Konfiguriert":', ' Nur Relais mit bereits eingetragenem Namen und Kategorie')

h3('Felder pro Relais')
simple_table(
    ['Feld', 'Beschreibung', 'Wertebereich'],
    [
        ['Gruppen-Nr.',  'Zugehörigkeit zu einer Gruppe. Relais mit gleicher Nummer schalten gemeinsam.', '0 = einzeln, 1–20 = Gruppe'],
        ['Name',         'Beschreibender Name des Fehlers, z.B. "PE unterbrochen"',                       'Freitext'],
        ['Kategorie',    'Messkategorie – welches Messverfahren findet diesen Fehler',                     'Auswahlliste (konfigurierbar)'],
        ['Stromkreis',   'Zuordnung zu einem der konfigurierten Stromkreise',                              'Auswahlliste (konfigurierbar)'],
    ]
)

h3('Gruppen')
body('Relais in derselben Gruppe (gleiche Gruppen-Nr. 1–20) schalten bei Aktivierung immer gemeinsam – sowohl im Prüfungsmodus als auch bei manueller Auswahl. Im Prüfungsmodus zählt eine Gruppe als einzelner Fehler.')

hint_box('Beispiel: Relais 5, 6 und 7 erhalten alle die Gruppen-Nr. 3. Wenn Relais 5 aktiviert wird, schalten auch 6 und 7 automatisch mit.', 'EFF3FF', '1F497D')

h3('Excel-Import und -Export')
numbered('Im Bereich "Excel Import/Export" eine leere Vorlage oder die aktuelle Konfiguration herunterladen.')
numbered('Die Datei lokal bearbeiten (Name, Gruppe, Kategorie, Stromkreis für jeden Relais eintragen).')
numbered('Die fertige Datei hochladen – das System übernimmt alle Werte.')

hint_box('Tipp: Es stehen auch vorgefertigte Templates für typische Messwand-Konfigurationen zur Auswahl.', 'FFF8E1', 'C07000')

h3('Statistik')
body('Oben auf der Seite wird eine Übersicht der Konfiguration angezeigt: Anzahl konfigurierter Relais, definierter Gruppen, Relais in Gruppen sowie Relais mit zugewiesener Kategorie.')
screenshot('07_relais_verwaltung.png', 'Relais-Verwaltung – Konfiguration aller 64 Relais')

# ── Konfiguration ─────────────────────────────────────────────────────────────
h2('Konfiguration')

h3('Stromkreise')
body('Stromkreise fassen mehrere Relais zu logischen Gruppen zusammen, die in der manuellen Fehlerauswahl und der Relais-Verwaltung als Einheit dargestellt werden.')

body('Vorhandene Systemstromkreise (nicht löschbar):')
simple_table(
    ['Nr.', 'Name', 'Relais'],
    [
        ['1', 'CEE 16A',               '0–9'],
        ['2', 'Herdanschlussdose',      '10–19'],
        ['3', 'Steckdose + Lampe 1',    '20–29'],
        ['4', 'Steckdose + Lampe 2',    '30–39'],
        ['5', 'Badsteckdose über RCBO', '40–49'],
        ['6', 'Wallbox',                '50–59'],
        ['7', 'RLO',                    '60–63'],
    ]
)

body('Eigene Stromkreise anlegen:')
numbered('Tab "Stromkreise" öffnen.')
numbered('"Neuer Stromkreis" drücken.')
numbered('Name, Beschreibung, Start-Relais-Nr. und Anzahl der Relais eingeben.')
numbered('Speichern.')

h3('Kategorien')
body('Kategorien ordnen Relais nach Messverfahren. Standard-Kategorien (nicht löschbar):')
simple_table(
    ['Kategorie', 'Bedeutung'],
    [
        ['RISO',     'Isolationswiderstandsmessung'],
        ['Zi',       'Netzinnenwiderstand'],
        ['Zs',       'Schleifenimpedanz'],
        ['Drehfeld', 'Drehfeldprüfung'],
        ['RCD',      'Fehlerstromschutzschalter-Test'],
    ]
)

body('Eigene Kategorie anlegen:')
numbered('Tab "Kategorien" öffnen.')
numbered('"Neue Kategorie" drücken.')
numbered('Namen eingeben und speichern.')
screenshot('08_konfiguration.png', 'Konfiguration – Stromkreise und Kategorien')

# ── Datenbank ─────────────────────────────────────────────────────────────────
h2('Datenbank')

body('Die Datenbank speichert alle durchgeführten Prüfungen automatisch. Jeder Eintrag enthält:')
bullet('Prüfungsnummer (z.B. VDE-12)')
bullet('Aktivierte Fehler (Namen der Relais)')
bullet('Datum und Uhrzeit')
bullet('Dauer in Sekunden')
bullet('Status (abgeschlossen oder abgebrochen)')

h3('Datenbank lesen')
numbered('Im Admin-Panel "Datenbank" öffnen.')
numbered('Die Tabelle zeigt alle Prüfungen in chronologischer Reihenfolge.')
numbered('Oben sind Statistiken eingeblendet: Gesamtanzahl, abgeschlossene und abgebrochene Prüfungen.')

h3('Datenbank exportieren (CSV)')
numbered('Auf der Datenbankseite "Als CSV exportieren" drücken.')
numbered('Die Datei wird heruntergeladen (Trennzeichen: Semikolon).')
numbered('Die CSV-Datei lässt sich mit Excel oder LibreOffice Calc öffnen.')

h3('Datenbank löschen')
numbered('"Datenbank leeren" drücken.')
numbered('Die Sicherheitsabfrage bestätigen – alle Einträge werden unwiderruflich gelöscht.')

hint_box('Achtung: Das Löschen der Datenbank ist nicht rückgängig zu machen. Vorher exportieren, wenn die Daten noch benötigt werden.', 'FFF0F0', 'C00000')
screenshot('09_datenbank.png', 'Datenbank – Prüfungsprotokoll und Exportfunktion')

# ── Netzwerk ──────────────────────────────────────────────────────────────────
h2('Netzwerk')

body('Die Netzwerkseite zeigt alle aktuellen Verbindungsadressen und ermöglicht die Verwaltung des Hotspots sowie die Einbindung in ein bestehendes Netzwerk.')

h3('Verbindungsadressen einsehen')
simple_table(
    ['Verbindungstyp', 'Adresse / Beispiel'],
    [
        ['Hostname',     'http://vde-messwand.local'],
        ['Ethernet-IP',  'http://192.168.1.x  (falls Kabel angeschlossen)'],
        ['WLAN-IP',      'http://192.168.1.y  (falls mit WLAN verbunden)'],
        ['Hotspot-IP',   'http://192.168.50.1  (wenn Hotspot aktiv)'],
    ]
)

h3('Hotspot verwalten')
bullet('SSID:', ' VDE-Messwand-1')
bullet('Passwort:', ' vde12345')
bullet('IP-Adresse:', ' 192.168.50.1')
body('Den Schalter im Bereich "Hotspot" betätigen, um ihn ein- oder auszuschalten.')

h3('WLAN-Netzwerk verbinden')
numbered('Im Bereich "WLAN" auf "Netzwerke suchen" drücken.')
numbered('Gewünschtes Netzwerk aus der Liste auswählen.')
numbered('Passwort eingeben und "Verbinden" drücken.')
numbered('Die neue IP-Adresse erscheint nach dem Verbinden in der Adressübersicht.')

hint_box('Hinweis: Wenn das System über WLAN und Hotspot gleichzeitig verbunden ist, sind beide Adressen erreichbar.', 'EFF3FF', '1F497D')
screenshot('10_netzwerk.png', 'Netzwerk – Verbindungsadressen und WLAN-Verwaltung')

# ── Übungsmodus-Konfiguration ─────────────────────────────────────────────────
h2('Übungsmodus-Konfiguration')

body('Hier wird festgelegt, welche Relais bei den Unterseiten des Übungsmodus für welche Messkategorie aktiviert werden. So können Auszubildende im Übungsmodus direkt passende Fehler schalten.')

h3('Konfiguration anlegen oder ändern')
numbered('Im Admin-Panel "Übungsmodus-Konfiguration" öffnen.')
numbered('Trainingsseite aus der Dropdown-Liste wählen (z.B. "Spannungsfrei" oder "Fluke").')
numbered('Kategorie auswählen (z.B. "RISO", "Zi").')
numbered('Die gewünschten Relais markieren und speichern.')

h3('Automatischer Import')
body('Über "Auto-Import aus Relais-Manager" werden alle Relais, die bereits in der Relais-Verwaltung mit einer Kategorie versehen sind, automatisch der passenden Trainingsseite zugeordnet.')

# ── Einstellungen ─────────────────────────────────────────────────────────────
h2('Einstellungen')

h3('Admin-Code ändern')
numbered('Im Admin-Panel "Einstellungen" öffnen.')
numbered('Den aktuellen Admin-Code eingeben (zur Verifikation).')
numbered('Den neuen Code eingeben (4–20 Ziffern).')
numbered('"Code ändern" drücken.')

h3('Wallbox-Verwaltung')
bullet('Wallbox installiert:', ' Wallbox-Stromkreis und Umschalter auf der Startseite sind sichtbar')
bullet('Wallbox nicht installiert:', ' Wallbox-Stromkreis ist ausgeblendet und wird nicht angesteuert')
screenshot('11_einstellungen.png', 'Einstellungen – Admin-Code und Wallbox')

# ── System ────────────────────────────────────────────────────────────────────
h2('Systemfunktionen')

h3('Relay Status Monitor')
body('Zeigt in Echtzeit den Zustand aller 64 Relais direkt vom Hardware-Interface. Grün = aktiv, Grau = inaktiv. Oben wird angezeigt, wie viele Relais aktuell aktiv sind.')
screenshot('12_relay_status.png', 'Relay Status Monitor – Live-Ansicht aller 64 Relais')

h3('Test-Durchlauf')
body('Prüft automatisch alle 64 Relais nacheinander: einschalten, per Modbus prüfen, ausschalten. Am Ende erscheint eine Zusammenfassung mit ggf. fehlerhaften Relais.')
numbered('Im Admin-Panel "Test-Durchlauf" öffnen.')
numbered('"Test starten" drücken.')
numbered('Fortschritt live verfolgen – am Ende: Ergebniszusammenfassung.')

hint_box('Hinweis: Der Test-Durchlauf dauert ca. 2–3 Minuten. Während des Tests können keine anderen Aktionen durchgeführt werden.', 'FFF8E1', 'C07000')
screenshot('13_test_mode.png', 'Test-Durchlauf – Automatischer Relais-Selbsttest')

h3('Alle Relais zurücksetzen')
body('Setzt alle 64 Relais sofort zurück. Nützlich, wenn eine Prüfung oder manuelle Auswahl abgebrochen werden soll.')
numbered('"Alle Relais zurücksetzen" drücken.')
numbered('Sicherheitsabfrage bestätigen.')

h3('System herunterfahren')
body('Fährt das Betriebssystem sicher herunter. Nach dem Herunterfahren ist das System erst nach einem manuellen Neustart wieder erreichbar.')
numbered('"System herunterfahren" drücken.')
numbered('Sicherheitsabfrage bestätigen. Alle Relais werden dabei zurückgesetzt.')

hint_box('Wichtig: Immer ordentlich herunterfahren, nicht einfach den Strom trennen. Nur so wird die SD-Karte des Raspberry Pi vor Datenverlust geschützt.', 'FFF0F0', 'C00000')

h3('Notaus-Erkennung')
body('Das System überwacht permanent zwei Hardware-Eingänge. Wenn ein Notaus-Kontakt betätigt wird, zeigt das System einen Warnhinweis ("NOTAUS BETÄTIGT") und fährt nach 120 Sekunden automatisch herunter, sofern der Kontakt nicht wieder geöffnet wird.')

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  ANHANG
# ══════════════════════════════════════════════════════════════════════════════

h1('Anhang – Kurzreferenz')

h2('Zugangsdaten und Adressen')
simple_table(
    ['Was', 'Wert'],
    [
        ['Hotspot SSID',       'VDE-Messwand-1'],
        ['Hotspot Passwort',   'vde12345'],
        ['Hotspot IP',         'http://192.168.50.1'],
        ['Hostname',           'http://vde-messwand.local'],
        ['Standard Admin-Code','1234'],
    ]
)

h2('Konfigurationsdateien')
simple_table(
    ['Datei', 'Inhalt'],
    [
        ['settings.json',      'Admin-Code, Wallbox-Einstellungen'],
        ['relais_config.json', 'Konfiguration aller 64 Relais (Name, Gruppe, Kategorie, Stromkreis)'],
        ['kategorien.json',    'Benutzerdefinierte Messkategorien'],
        ['vde_messwand.db',    'SQLite-Datenbank mit allen Prüfungsprotokollen'],
        ['hotspot_state.json', 'Aktueller Status des Hotspots'],
    ]
)

h2('Häufige Fragen')

h3('Kein Zugriff auf die Oberfläche?')
numbered('Prüfen, ob das Gerät eingeschaltet ist (ca. 60 Sekunden Startzeit einplanen).')
numbered('Mit WLAN "VDE-Messwand-1" verbinden, Passwort: vde12345.')
numbered('Browser öffnen und http://192.168.50.1 eingeben.')
numbered('Falls kein Hotspot: Im Admin-Bereich Netzwerkstatus prüfen.')

h3('Admin-Code vergessen?')
body('Den Code kann direkt in der Datei settings.json auf dem Gerät zurückgesetzt werden. Dazu ist ein direkter Zugriff per SSH oder über das angeschlossene Display erforderlich.')

h3('Relais reagiert nicht?')
numbered('Test-Durchlauf im Admin-Bereich starten.')
numbered('Relay Status Monitor öffnen und prüfen, ob der Relay-Zustand sich ändert.')
numbered('Modbus-Verbindung prüfen (Kabel zwischen Raspberry Pi und Relaismodulen).')

# ── Speichern ─────────────────────────────────────────────────────────────────
output = '/home/vde/VDE-Messwand/Handbuch_VDE-Messwand.docx'
doc.save(output)
print(f'Gespeichert: {output}')
